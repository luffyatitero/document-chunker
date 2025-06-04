import os
import magic
import aiofiles
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import PyPDF2
import docx
import pandas as pd
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    CharacterTextSplitter,
    TokenTextSplitter,
)
import tiktoken
import chardet
from io import BytesIO

# Import enums and config from schemas for type safety and alignment
from app.schemas.document import SplitterType, LengthFunction, ChunkingConfigBase

class DocumentProcessor:
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'application/pdf',
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.xls': 'application/vnd.ms-excel',
        '.csv': 'text/csv',
        '.html': 'text/html',
        '.htm': 'text/html',
        '.py': 'text/x-python',
        '.js': 'text/javascript',
        '.java': 'text/x-java-source',
        '.cpp': 'text/x-c++src',
        '.c': 'text/x-csrc',
        '.json': 'application/json',
        '.xml': 'application/xml',
        '.rtf': 'application/rtf'
    }

    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)

    async def extract_text_from_file(self, file_path: str, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text content from various file types"""
        metadata = {}
        try:
            if file_type == 'application/pdf':
                return await self._extract_from_pdf(file_path, metadata)
            elif file_type in [
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/msword'
            ]:
                return await self._extract_from_docx(file_path, metadata)
            elif file_type in [
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                'application/vnd.ms-excel'
            ]:
                return await self._extract_from_excel(file_path, metadata)
            elif file_type == 'text/csv':
                return await self._extract_from_csv(file_path, metadata)
            elif 'text/' in file_type or file_type in ['application/json', 'application/xml']:
                return await self._extract_from_text(file_path, metadata)
            else:
                # Try to extract as text for unknown types
                return await self._extract_from_text(file_path, metadata)
        except Exception as e:
            raise ValueError(f"Failed to extract text from file: {str(e)}")

    async def _extract_from_pdf(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        async with aiofiles.open(file_path, 'rb') as file:
            content = await file.read()
            pdf_reader = PyPDF2.PdfReader(BytesIO(content))
            text = ""
            metadata['pages'] = len(pdf_reader.pages)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text() or ""
                text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            return text.strip(), metadata

    async def _extract_from_docx(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            raise ValueError(f"Failed to read DOCX file: {str(e)}")
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        metadata['paragraphs'] = len(doc.paragraphs)
        return text.strip(), metadata

    async def _extract_from_excel(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        try:
            df = pd.read_excel(file_path, sheet_name=None)  # Read all sheets
        except Exception as e:
            raise ValueError(f"Failed to read Excel file: {str(e)}")
        text = ""
        for sheet_name, sheet_df in df.items():
            text += f"\n--- Sheet: {sheet_name} ---\n"
            # Limit rows for very large sheets
            if len(sheet_df) > 1000:
                text += sheet_df.head(1000).to_string(index=False) + "\n... (truncated)\n"
            else:
                text += sheet_df.to_string(index=False) + "\n"
        metadata['sheets'] = list(df.keys())
        metadata['total_rows'] = sum(len(sheet_df) for sheet_df in df.values())
        return text.strip(), metadata

    async def _extract_from_csv(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        try:
            df = pd.read_csv(file_path)
        except Exception as e:
            raise ValueError(f"Failed to read CSV file: {str(e)}")
        # Limit rows for very large CSVs
        if len(df) > 1000:
            text = df.head(1000).to_string(index=False) + "\n... (truncated)"
        else:
            text = df.to_string(index=False)
        metadata['rows'] = len(df)
        metadata['columns'] = list(df.columns)
        return text, metadata

    async def _extract_from_text(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        # Detect encoding
        async with aiofiles.open(file_path, 'rb') as file:
            raw_data = await file.read()
            encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
        async with aiofiles.open(file_path, 'r', encoding=encoding) as file:
            text = await file.read()
        metadata['encoding'] = encoding
        metadata['lines'] = len(text.split('\n'))
        return text, metadata

    def get_text_splitter(self, config: ChunkingConfigBase):
        """Create and configure text splitter based on config"""
        # Use enums for type safety
        splitter_type = config.splitter_type or SplitterType.RECURSIVE
        length_function = self._get_length_function(
            config.length_function.value if hasattr(config.length_function, "value") else config.length_function
        )
        # Default values for advanced options
        is_separator_regex = getattr(config, "is_separator_regex", False)
        keep_separator = getattr(config, "keep_separator", False)
        add_start_index = getattr(config, "add_start_index", False)
        strip_whitespace = getattr(config, "strip_whitespace", True)
        separators = getattr(config, "custom_separators", None) or getattr(config, "separators", None)

        if splitter_type == SplitterType.RECURSIVE:
            return RecursiveCharacterTextSplitter(
                chunk_size=config.chunk_size,
                chunk_overlap=config.chunk_overlap,
                separators=separators or ["\n\n", "\n", " ", ""],
                length_function=length_function,
                is_separator_regex=is_separator_regex,
                keep_separator=keep_separator,
                add_start_index=add_start_index,
                strip_whitespace=strip_whitespace
            )
        elif splitter_type == SplitterType.CHARACTER:
            separator = (separators[0] if separators else "\n")
            return CharacterTextSplitter(
                chunk_size=config.chunk_size,
                chunk_overlap=config.chunk_overlap,
                separator=separator,
                length_function=length_function,
                is_separator_regex=is_separator_regex,
                keep_separator=keep_separator,
                add_start_index=add_start_index,
                strip_whitespace=strip_whitespace
            )
        elif splitter_type == SplitterType.TOKEN:
            return TokenTextSplitter(
                chunk_size=config.chunk_size,
                chunk_overlap=config.chunk_overlap,
                encoding_name="cl100k_base",
                add_start_index=add_start_index
            )
        else:
            # Warn and fallback
            import warnings
            warnings.warn(f"Unknown splitter_type '{splitter_type}', defaulting to RecursiveCharacterTextSplitter.")
            return RecursiveCharacterTextSplitter(
                chunk_size=config.chunk_size,
                chunk_overlap=config.chunk_overlap
            )

    def _get_length_function(self, length_func: Optional[str]):
        """Get the appropriate length function"""
        if length_func == LengthFunction.TIKTOKEN or length_func == "tiktoken":
            encoder = tiktoken.get_encoding("cl100k_base")
            return lambda text: len(encoder.encode(text))
        else:
            return len